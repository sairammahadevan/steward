"""
Document parser for STEWARD.

Handles three document types:
- Text PDF       — has selectable text, extracted directly
- Scanned PDF    — image-based pages, converted to PNG for Claude vision
- DOCX           — extracted via python-docx
- Plain text     — passed through as-is

Detection: if a PDF page yields < 50 characters, it's treated as scanned.
"""

from __future__ import annotations

import hashlib
import logging
import tempfile
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Union

logger = logging.getLogger(__name__)

SCANNED_TEXT_THRESHOLD = 50  # chars per page below this = scanned
IMAGE_DPI = 100              # lower DPI = smaller images = faster API calls (still readable)
MAX_SCANNED_PAGES = 8        # cap pages sent to Claude — enough to identify any document


class DocType(str, Enum):
    TEXT_PDF = "text_pdf"
    SCANNED_PDF = "scanned_pdf"
    MIXED_PDF = "mixed_pdf"
    DOCX = "docx"
    TEXT = "text"


@dataclass
class ParsedDocument:
    """Result of parsing a document."""
    doc_type: DocType
    sha256: str
    filename: str
    text: str = ""                        # populated for text-based docs
    image_paths: list[Path] = field(default_factory=list)  # populated for scanned docs
    page_count: int = 0
    error: str = ""

    @property
    def is_vision(self) -> bool:
        """True if this document needs to be sent as images to Claude."""
        return self.doc_type in (DocType.SCANNED_PDF, DocType.MIXED_PDF)

    @property
    def success(self) -> bool:
        return not self.error


def compute_sha256(path: Path) -> str:
    """Compute SHA-256 hash of a file."""
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def parse_document(
    path: Path,
    tmp_dir: Path | None = None,
    image_dpi: int = IMAGE_DPI,
) -> ParsedDocument:
    """
    Parse a document file and return a ParsedDocument.

    Args:
        path:      Absolute path to the document.
        tmp_dir:   Directory to write page images into (for scanned PDFs).
                   If None, a system temp dir is used.
        image_dpi: DPI for rasterising scanned pages. Use 150 for higher-quality
                   re-ingests; leave at default (100) for first-time ingests.

    Returns:
        ParsedDocument with text or image_paths populated.
    """
    filename = path.name
    sha256 = compute_sha256(path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _parse_pdf(path, filename, sha256, tmp_dir, image_dpi=image_dpi)
    elif suffix in (".docx", ".doc"):
        return _parse_docx(path, filename, sha256)
    elif suffix in (".txt", ".md"):
        return _parse_text(path, filename, sha256)
    else:
        return ParsedDocument(
            doc_type=DocType.TEXT,
            sha256=sha256,
            filename=filename,
            error=f"Unsupported file type: {suffix}",
        )


def parse_text_content(text: str, title: str) -> ParsedDocument:
    """Parse plain text content (pasted by user)."""
    sha256 = hashlib.sha256(text.encode()).hexdigest()
    return ParsedDocument(
        doc_type=DocType.TEXT,
        sha256=sha256,
        filename=title or "pasted-text",
        text=text.strip(),
        page_count=1,
    )


# ── PDF ──────────────────────────────────────────────────────────────────────

def _parse_pdf(path: Path, filename: str, sha256: str, tmp_dir: Path | None, image_dpi: int = IMAGE_DPI) -> ParsedDocument:
    """Parse a PDF — detect text vs scanned and route accordingly."""
    try:
        import fitz  # pymupdf
    except ImportError:
        return ParsedDocument(
            doc_type=DocType.TEXT_PDF,
            sha256=sha256,
            filename=filename,
            error="pymupdf not installed. Run: pip install pymupdf",
        )

    try:
        doc = fitz.open(str(path))
    except Exception as e:
        return ParsedDocument(
            doc_type=DocType.TEXT_PDF,
            sha256=sha256,
            filename=filename,
            error=f"Could not open PDF: {e}",
        )

    text_pages: list[str] = []
    scanned_pages: list[int] = []

    for i, page in enumerate(doc):
        page_text = page.get_text().strip()
        if len(page_text) >= SCANNED_TEXT_THRESHOLD:
            text_pages.append(page_text)
        else:
            scanned_pages.append(i)
            text_pages.append("")  # placeholder

    total_pages = len(doc)

    # Determine document type
    if not scanned_pages:
        doc_type = DocType.TEXT_PDF
    elif len(scanned_pages) == total_pages:
        doc_type = DocType.SCANNED_PDF
    else:
        doc_type = DocType.MIXED_PDF

    logger.info(
        "%s: %d pages, %d scanned, type=%s",
        filename, total_pages, len(scanned_pages), doc_type,
    )

    # For text/mixed: join all text
    full_text = "\n\n".join(t for t in text_pages if t)

    # For scanned/mixed: render scanned pages to images (capped at MAX_SCANNED_PAGES)
    image_paths: list[Path] = []
    if scanned_pages:
        out_dir = tmp_dir or Path(tempfile.mkdtemp())
        out_dir.mkdir(parents=True, exist_ok=True)

        pages_to_render = scanned_pages[:MAX_SCANNED_PAGES]
        if len(scanned_pages) > MAX_SCANNED_PAGES:
            logger.info(
                "%s: %d scanned pages, capping at %d for API call",
                filename, len(scanned_pages), MAX_SCANNED_PAGES,
            )

        for page_idx in pages_to_render:
            page = doc[page_idx]
            mat = fitz.Matrix(image_dpi / 72, image_dpi / 72)
            pix = page.get_pixmap(matrix=mat)
            img_path = out_dir / f"{sha256}_page_{page_idx}.png"
            pix.save(str(img_path))
            image_paths.append(img_path)
            logger.debug("Rendered page %d → %s", page_idx, img_path)

    doc.close()

    return ParsedDocument(
        doc_type=doc_type,
        sha256=sha256,
        filename=filename,
        text=full_text,
        image_paths=image_paths,
        page_count=total_pages,
    )


# ── DOCX ─────────────────────────────────────────────────────────────────────

def _parse_docx(path: Path, filename: str, sha256: str) -> ParsedDocument:
    """Extract text from a DOCX file."""
    try:
        from docx import Document
    except ImportError:
        return ParsedDocument(
            doc_type=DocType.DOCX,
            sha256=sha256,
            filename=filename,
            error="python-docx not installed. Run: pip install python-docx",
        )

    try:
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        full_text = "\n\n".join(paragraphs)
        return ParsedDocument(
            doc_type=DocType.DOCX,
            sha256=sha256,
            filename=filename,
            text=full_text,
            page_count=1,
        )
    except Exception as e:
        return ParsedDocument(
            doc_type=DocType.DOCX,
            sha256=sha256,
            filename=filename,
            error=f"Could not read DOCX: {e}",
        )


# ── Plain text ────────────────────────────────────────────────────────────────

def _parse_text(path: Path, filename: str, sha256: str) -> ParsedDocument:
    """Read a plain text or markdown file."""
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
        return ParsedDocument(
            doc_type=DocType.TEXT,
            sha256=sha256,
            filename=filename,
            text=text.strip(),
            page_count=1,
        )
    except Exception as e:
        return ParsedDocument(
            doc_type=DocType.TEXT,
            sha256=sha256,
            filename=filename,
            error=f"Could not read file: {e}",
        )
